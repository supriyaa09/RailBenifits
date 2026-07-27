import os
import sys
import docx
import fitz  # PyMuPDF
import tempfile
from unittest.mock import MagicMock

def parse_docx(file_path):
    doc = docx.Document(file_path)
    text_content = []
    # Read paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)
    # Read tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(" | ".join(row_text))
    return "\n".join(text_content)

def parse_pdf_hybrid(file_path):
    doc = fitz.open(file_path)
    total_pages = len(doc)
    
    pages_requiring_ocr = []
    text_content = []
    chars_per_page = []
    ocr_confidences = []
    
    ocr = None  # Lazy initialize OCR only when required
    
    for page_num in range(total_pages):
        page = doc[page_num]
        page_text = page.get_text()
        
        # If searchable text is very short/empty, page is considered scanned and triggers OCR
        if len(page_text.strip()) < 100:
            pages_requiring_ocr.append(page_num)
            
            # Lazy initialize PaddleOCR with environment mitigations
            if ocr is None:
                sys.stderr.write("[Scanned page detected. Lazy-initializing PaddleOCR pipeline...]\n")
                sys.stderr.flush()
                
                # Mock modelscope to avoid loading PyTorch and hitting shm.dll loading library bugs
                sys.modules['modelscope'] = MagicMock()
                
                # Hook create_predictor to disable CPU static executor oneDNN instruction crashes on Windows
                import paddle.inference as paddle_inference
                if hasattr(paddle_inference, 'create_predictor'):
                    old_cp = paddle_inference.create_predictor
                    paddle_inference.create_predictor = lambda c: (
                        c.disable_onednn() if hasattr(c, 'disable_onednn') else None,
                        c.disable_onednn_fc_passes() if hasattr(c, 'disable_onednn_fc_passes') else None,
                        old_cp(c)
                    )[-1]
                
                from paddleocr import PaddleOCR
                ocr = PaddleOCR(use_textline_orientation=True, lang='en')
            
            # Render page to a temporary image
            pix = page.get_pixmap(dpi=200)
            temp_dir = tempfile.gettempdir()
            temp_img_path = os.path.join(temp_dir, f"ocr_page_{os.getpid()}_{page_num}.png")
            pix.save(temp_img_path)
            
            page_ocr_text = []
            page_ocr_scores = []
            
            try:
                # Run OCR predict
                results = ocr.predict(temp_img_path)
                if results:
                    for res in results:
                        # OCRResult is a dict-like object — use key access, not attribute access
                        texts = res['rec_texts'] if 'rec_texts' in res else []
                        scores = res['rec_scores'] if 'rec_scores' in res else []
                        if texts:
                            page_ocr_text.extend(texts)
                        if scores:
                            page_ocr_scores.extend(scores)
                        sys.stderr.write(f"[Page {page_num}: {len(texts)} text regions detected]\n")
                        sys.stderr.flush()
            except Exception as ocr_err:
                sys.stderr.write(f"[Page {page_num} OCR Error: {str(ocr_err)}]\n")
                sys.stderr.flush()
            finally:
                if os.path.exists(temp_img_path):
                    try:
                        os.remove(temp_img_path)
                    except Exception:
                        pass
            
            extracted_page_text = "\n".join(page_ocr_text)
            text_content.append(extracted_page_text)
            chars_per_page.append(len(extracted_page_text))
            
            if page_ocr_scores:
                ocr_confidences.append(sum(page_ocr_scores) / len(page_ocr_scores))
            else:
                ocr_confidences.append(0.0)
        else:
            # Searchable page, skip OCR
            text_content.append(page_text)
            chars_per_page.append(len(page_text))
            ocr_confidences.append(1.0) # 100% confidence for searchable PDF text
            
    # Print statistics logs to stderr
    avg_ocr_confidence = (sum(ocr_confidences) / len(ocr_confidences)) if ocr_confidences else 1.0
    sys.stderr.write(f"--- OCR Processing Stats ---\n")
    sys.stderr.write(f"Number of pages processed: {total_pages}\n")
    sys.stderr.write(f"Pages requiring OCR: {pages_requiring_ocr}\n")
    sys.stderr.write(f"Characters extracted per page: {chars_per_page}\n")
    sys.stderr.write(f"OCR confidence: {avg_ocr_confidence:.4f}\n")
    sys.stderr.write(f"----------------------------\n")
    sys.stderr.flush()
    
    return "\n".join(text_content)

def parse_txt(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def main():
    if len(sys.argv) < 2:
        print("Usage: python parse_circular.py <file_path>")
        sys.exit(1)
        
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
        
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == ".docx":
            text = parse_docx(file_path)
        elif ext == ".pdf":
            text = parse_pdf_hybrid(file_path)
        elif ext in [".txt", ".md"]:
            text = parse_txt(file_path)
        else:
            text = parse_txt(file_path)
            
        # Clean and normalize text: replace multiple spaces with single spaces, clean up carriage returns
        cleaned_lines = []
        for line in text.split("\n"):
            line_str = " ".join(line.split())
            if line_str:
                cleaned_lines.append(line_str)
        normalized_text = "\n".join(cleaned_lines)

        # Write to stdout using UTF-8 to prevent console encoding issues
        sys.stdout.buffer.write(normalized_text.encode("utf-8"))
    except Exception as e:
        print(f"Error parsing file: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
